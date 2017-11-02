import os.path
from docx import Document
import code
import pickle
import re
import transform
from  docx.oxml.ns import  qn

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)

with open('suplist.pkl','rb') as f:
	SUPdic = pickle.load(f)

temppath = r'C:\Users\lenovo\Desktop\Contract\硬件采购合同模板1.2.docx'
stemppath = r'C:\Users\lenovo\Desktop\Contract\软件采购合同模板1.2.docx'

dir = r'C:\Users\lenovo\Desktop\待办呈批件\硬件采购'
sdir = r'C:\Users\lenovo\Desktop\待办呈批件\软件采购'


def hardwareauto():
	dirlist = os.listdir(dir)
	dirlist = [x for x in dirlist if '.py' not in x]
	lenlist = len(dirlist)
	
	SUPcount = {}
	figure = []
	
	for i in range(0,lenlist):
		path = os.path.join(dir,dirlist[i])
		filename = '666'+ dirlist[i]
		document = Document(path)
		try:
			table = document.tables[0]
		except IndexError:
			pass
		print(path)
		 
		for key in SUPdic.keys():
			for para in document.paragraphs[7:]:
				if key in para.text:
					SUPcount[key] = 0
			for para in document.paragraphs[7:]:
				if key in para.text:
					SUPcount[key] += 1
					supname = max(SUPcount, key=SUPcount.get)
		if len(SUPcount) < 3:
			new = 1
		else:
			new = 0
		for para in document.paragraphs[7:]:
			fig = re.findall(r'\d*\,*\d+,\d+\.*\d*',para.text)
	
			if (any(fig) and type(fig)==list and len(fig) == 1):
				figure.append(fig)
			f_figure = sum(figure,[])#turn list of lists to a list
	
	
		for i in range(len(f_figure)):
			f_figure[i] = float(f_figure[i].replace(',',''))
		amount = min(f_figure)
		pretax = round(amount/1.17,2)
		tax = round(amount - pretax,2)
	
		if __name__ == "__main__":
			CNamount = transform.cncurrency(amount)
			CNpretax = transform.cncurrency(pretax)
			CNtax = transform.cncurrency(tax)
	
		famount = format(amount,',')
		fpretax = format(pretax,',')
		ftax = format(tax,',')
	
		SUPcount.clear()
		figure=[]
	
		template = Document(temppath)
		template.paragraphs[4].add_run(supname+'（下称乙方）')
		template.paragraphs[6].runs[1].add_text(supname)
		template.paragraphs[9].add_run(famount + '（小写）')
		template.paragraphs[10].add_run(CNamount)
		template.paragraphs[12].add_run(fpretax + '（小写）')
		template.paragraphs[13].add_run(CNpretax)
		template.paragraphs[14].add_run(ftax+ '（小写）')
		template.paragraphs[15].add_run(CNtax)
		template.paragraphs[17].add_run(famount+'元（大写人民币:'+CNamount+'）')
		for i in [9,10,12,13,14,15,17]:
			template.paragraphs[i].runs[-1].underline = True
		#It doesn't recognize '华文楷体'？？？
		fontname = template.paragraphs[6].runs[1].font.name
		for i in [4,9,10,12,13,14,15,17]:
			template.paragraphs[i].runs[-1].font.name = fontname
			template.paragraphs[i].runs[-1]._element.rPr.rFonts.set(qn('w:eastAsia'), '华文楷体')
			template.paragraphs[i].runs[-1].font.size = 177800
		template.paragraphs[6].runs[1].font.name = fontname
		template.paragraphs[6].runs[1]._element.rPr.rFonts.set(qn('w:eastAsia'), '华文楷体')
		template.paragraphs[6].runs[1].font.size = 177800
	
		move_table_after(table,template.paragraphs[20])
	
	
		for i in range(len(template.tables[-1].column_cells(1))):
			template.tables[-1].column_cells(1)[i].text = SUPdic[supname][i]
	
		if new == 1:
			print('supplier is new!')
			template.save(os.path.join(dir,'new'+ filename))
		else:
			template.save(os.path.join(dir,filename))
		

def softwareauto():
	sdirlist = os.listdir(sdir)
	lenlist = len(sdirlist)
	
	SUPcount = {}
	figure = []
	
	for i in range(0,lenlist):
		path = os.path.join(sdir,sdirlist[i])
		filename = '666'+ sdirlist[i]
		document = Document(path)
		try:
			table = document.tables[0]
		except IndexError:
			pass
		print(path)
	
		for key in SUPdic.keys():
			for para in document.paragraphs[7:]:
				if key in para.text:
					SUPcount[key] = 0
			for para in document.paragraphs[7:]:
				if key in para.text:
					SUPcount[key] += 1
					supname = max(SUPcount, key=SUPcount.get)
		if len(SUPcount) < 3:
			new = 1
		else:
			new = 0		
		for para in document.paragraphs[7:]:
			fig = re.findall(r'\d+,\d+\.*\d*',para.text)
	
			if (any(fig) and type(fig)==list and len(fig) == 1):
				figure.append(fig)
			f_figure = sum(figure,[])#turn list of lists to a list
	
	
		for i in range(len(f_figure)):
			f_figure[i] = float(f_figure[i].replace(',',''))
		amount = min(f_figure)
		pretax = round(amount/1.17,2)
		tax = round(amount - pretax,2)
	
		if __name__ == "__main__":
			CNamount = transform.cncurrency(amount)
			CNpretax = transform.cncurrency(pretax)
			CNtax = transform.cncurrency(tax)
	
		famount = format(amount,',')
		fpretax = format(pretax,',')
		ftax = format(tax,',')
	
		SUPcount.clear()
		figure=[]
	
		template = Document(stemppath)
		template.paragraphs[4].add_run(supname+'（下称乙方）')
		template.paragraphs[6].runs[1].add_text(supname)
		template.paragraphs[9].add_run(famount + '（小写）')
		template.paragraphs[10].add_run(CNamount)
		template.paragraphs[12].add_run(fpretax + '（小写）')
		template.paragraphs[13].add_run(CNpretax)
		template.paragraphs[14].add_run(ftax+ '（小写）')
		template.paragraphs[15].add_run(CNtax)
		template.paragraphs[17].add_run(famount+'元（大写人民币:'+CNamount+'）')
		for i in [9,10,12,13,14,15,17]:
			template.paragraphs[i].runs[-1].underline = True
		fontname = template.paragraphs[6].runs[1].font.name
		for i in [4,9,10,12,13,14,15,17]:
			template.paragraphs[i].runs[-1].font.name = fontname
			template.paragraphs[i].runs[-1]._element.rPr.rFonts.set(qn('w:eastAsia'), '华文楷体')
			template.paragraphs[i].runs[-1].font.size = 177800
		template.paragraphs[6].runs[1].font.name = fontname
		template.paragraphs[6].runs[1]._element.rPr.rFonts.set(qn('w:eastAsia'), '华文楷体')
		template.paragraphs[6].runs[1].font.size = 177800
	
		move_table_after(table,template.paragraphs[18])
	
	
		for i in range(len(template.tables[-1].column_cells(1))):
			template.tables[-1].column_cells(1)[i].text = SUPdic[supname][i]
			
		if new == 1:
			print('shit is new!')
			template.save(os.path.join(sdir,'new'+ filename))
		else:
			template.save(os.path.join(sdir,filename))

selection = input('>>合同类别：')
if selection == '硬':
	hardwareauto()
if selection == '软':
	softwareauto()

#code.interact(banner = "", local = locals())
#for i in range(0,len(document.paragraphs)):
#	print(str(i)+'  ' + document.paragraphs[i].text)
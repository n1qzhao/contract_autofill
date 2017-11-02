import os.path
from docx import Document
import pickle

SUP_info_dic = {} #save all the suppliers' info here
 

def get_cells_txt(cell):
	return cell.text 

dir = r'C:\your contracts directory'
dirlist = os.listdir(dir)
docxlist = []

for filename in dirlist:
	if '.docx' in filename:
		docxlist.append(filename)


for i in range(0,len(docxlist)):
	path = os.path.join(dir,docxlist[i])
	document = Document(path)
	try:
		sup_name = document.paragraphs[4].text.strip(' ')[3:-6]
	except IndexError:
		pass
	
	for table in document.tables:
		try:
			if table.column_cells(1)[1].text == '法定代表人：':
				content = list(map(get_cells_txt,table.column_cells(1)))
				
				if sup_name in SUP_info_dic.keys():
					print(content)
					call = input('Info has newer version, type OK to replace?')
					if call == 'OK':
						SUP_info_dic[sup_name] = content	
					else:
						pass
				else:
					SUP_info_dic[sup_name] = content
		except IndexError:
			print('.')

SUP_info_dic.pop['']

with open('suplist.pkl','wb') as f:
	pickle.dump(SUP_info_dic,f)
f.close()			
